from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
import os
import json

# Function to add hyperlink to a paragraph
def add_hyperlink(paragraph, url, text, font_size=Pt(9), color=RGBColor(0, 0, 255)):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set font size
    sz = OxmlElement('w:sz')
    sz.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', str(int(font_size.pt * 2)))
    rPr.append(sz)
    
    # Set color
    c = OxmlElement('w:color')
    c.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(c)
    
    # Set underline
    u = OxmlElement('w:u')
    u.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

# --- LOAD DATA FROM JSON FILES ---
def load_json(filename):
    filepath = os.path.join('data', filename)
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)

# Load all data files
personal_info = load_json('personal-info.json')
summary_data = load_json('summary.json')
skills_data = load_json('skills.json')
experience_data = load_json('experience.json')
education_data = load_json('education.json')
certifications_data = load_json('certifications.json')
section_labels = load_json('section-labels.json')

# Initialize Document
doc = Document()

# --- PAGE LAYOUT SETTINGS ---
# Set margins to 0.5 inches for all sides
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# --- DOCUMENT HEADER ---
# Enable different first page header and use first_page_header for first page only
for section in sections:
    section.different_first_page_header_footer = True

# Access the first page header (appears only on first page)
header = sections[0].first_page_header
header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

# Create 3-column table in header: Email/Phone | Name | LinkedIn/GitHub
header_table = header.add_table(rows=1, cols=3, width=Inches(7.5))
header_table.allow_autofit = False
header_table.columns[0].width = Inches(2.0)
header_table.columns[1].width = Inches(3.5)
header_table.columns[2].width = Inches(2.0)

# Set cell widths directly to ensure they're respected
header_table.cell(0, 0).width = Inches(2.0)
header_table.cell(0, 1).width = Inches(3.5)
header_table.cell(0, 2).width = Inches(2.0)

# Left column: Email and Phone (stacked)
left_cell = header_table.cell(0, 0)
left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
left_p = left_cell.paragraphs[0]
left_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
add_hyperlink(left_p, f"mailto:{personal_info['email']}", personal_info['email'], font_size=Pt(8), color=RGBColor(0, 102, 204))
left_p.add_run('\n')
phone_run = left_p.add_run(personal_info['phone'])
phone_run.font.size = Pt(8)

# Center column: Name
center_cell = header_table.cell(0, 1)
center_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
center_p = center_cell.paragraphs[0]
center_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_run = center_p.add_run(personal_info['name'].upper())
name_run.bold = True
name_run.font.size = Pt(18)
name_run.font.color.rgb = RGBColor(46, 64, 83)

# Right column: LinkedIn and GitHub (horizontal)
right_cell = header_table.cell(0, 2)
right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
right_p = right_cell.paragraphs[0]
right_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
add_hyperlink(right_p, personal_info['linkedin'], 'LinkedIn', font_size=Pt(8), color=RGBColor(0, 102, 204))
right_p.add_run(' | ').font.size = Pt(8)
add_hyperlink(right_p, personal_info['github'], 'GitHub', font_size=Pt(8), color=RGBColor(0, 102, 204))

# Remove the empty first paragraph in header to reduce space
if header_para.text == '':
    p_element = header_para._element
    p_element.getparent().remove(p_element)

# --- STYLES & FORMATTING ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(8)

# Function to add section header (without background)
def add_section_header(document, text):
    p = document.add_paragraph()
    runner = p.add_run(text.upper())
    runner.bold = True
    runner.font.size = Pt(9)
    runner.font.color.rgb = RGBColor(46, 64, 83)  # Dark blue text

    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)

def add_role_header(document, title, company_location, date):
    # Table for layout: Title (Left) | Date (Right)
    table = document.add_table(rows=1, cols=2)
    table.allow_autofit = False

    # Set table width to full available width (8.5" - 0.5" - 0.5" = 7.5")
    table.width = Inches(7.5)

    # Set column widths with proper cell width assignment
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(2.0)

    # Also set the cell widths directly to ensure they're respected
    table.cell(0, 0).width = Inches(5.5)
    table.cell(0, 1).width = Inches(2.0)

    # Title Cell
    cell_1 = table.cell(0, 0)
    p1 = cell_1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(46, 64, 83) # Dark Blue/Grey

    # Date Cell
    cell_2 = table.cell(0, 1)

    # Set cell vertical alignment
    cell_2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p2 = cell_2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r2 = p2.add_run(date)
    r2.bold = True
    r2.font.size = Pt(9)

    # Company Line below title (only if provided - for backward compatibility)
    if company_location:
        p_sub = document.add_paragraph()
        r_sub = p_sub.add_run(company_location)
        r_sub.italic = True
        p_sub.paragraph_format.space_after = Pt(1)

# --- PROFESSIONAL SUMMARY ---
add_section_header(doc, section_labels['professional_summary'])
summary = doc.add_paragraph(summary_data['text'])
summary.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
summary.paragraph_format.space_after = Pt(0)

# --- KEY SKILLS ---
add_section_header(doc, section_labels['key_skills'])

for skill_category in skills_data['categories']:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    cat_run = p.add_run(skill_category['category'] + " ")
    cat_run.bold = True
    p.add_run(skill_category['items'])

# --- PROFESSIONAL EXPERIENCE ---
add_section_header(doc, section_labels['professional_experience'])

for company in experience_data['companies']:
    # Add company header
    company_p = doc.add_paragraph()
    company_run = company_p.add_run(f"{company['company']} | {company['location']}")
    company_run.bold = True
    company_run.font.size = Pt(9)
    company_run.font.color.rgb = RGBColor(0, 51, 102)  # Darker blue
    company_p.paragraph_format.space_before = Pt(2)
    company_p.paragraph_format.space_after = Pt(0)

    # Add roles under this company
    for role in company['roles']:
        add_role_header(doc, role['title'], "", role['date'])
        for bullet in role['bullets']:
            bullet_p = doc.add_paragraph(bullet, style='List Bullet')
            bullet_p.paragraph_format.space_after = Pt(0)
            bullet_p.paragraph_format.space_before = Pt(0)
            bullet_p.paragraph_format.line_spacing = 1.0
            bullet_p.paragraph_format.left_indent = Inches(0.25)
            bullet_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in bullet_p.runs:
                run.font.size = Pt(8)

# Earlier Career
add_section_header(doc, section_labels['earlier_career'])

# Additional Companies (Earlier Career Details)
if 'additional_companies' in experience_data:
    for company in experience_data['additional_companies']:
        # Add company header
        company_p = doc.add_paragraph()
        company_run = company_p.add_run(f"{company['company']} | {company['location']}")
        company_run.bold = True
        company_run.font.size = Pt(9)
        company_run.font.color.rgb = RGBColor(0, 51, 102)  # Darker blue
        company_p.paragraph_format.space_before = Pt(2)
        company_p.paragraph_format.space_after = Pt(0)

        # Add roles under this company
        for role in company['roles']:
            add_role_header(doc, role['title'], "", role['date'])
            for bullet in role['bullets']:
                bullet_p = doc.add_paragraph(bullet, style='List Bullet')
                bullet_p.paragraph_format.space_after = Pt(0)
                bullet_p.paragraph_format.space_before = Pt(0)
                bullet_p.paragraph_format.line_spacing = 1.0
                bullet_p.paragraph_format.left_indent = Inches(0.25)
                bullet_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in bullet_p.runs:
                    run.font.size = Pt(8)

# --- EDUCATION ---
add_section_header(doc, section_labels['education'])
for degree in education_data['degrees']:
    edu_p = doc.add_paragraph()
    edu_p.paragraph_format.space_after = Pt(0)
    edu_p.paragraph_format.space_before = Pt(0)

    # Degree name in bold
    degree_run = edu_p.add_run(degree['degree'])
    degree_run.bold = True

    # Institution in regular text on same line
    edu_p.add_run(f", {degree['institution']}")

# --- CERTIFICATIONS ---
add_section_header(doc, section_labels['certifications'])

# Create a 2-column table for certifications (compact layout)
cert_list = certifications_data['certifications']
# Split each certification line by ' | ' to get individual certs
all_certs = []
for cert_line in cert_list:
    all_certs.extend([c.strip() for c in cert_line.split(' | ')])

# Create table with 2 columns
num_rows = (len(all_certs) + 1) // 2  # Calculate rows needed
cert_table = doc.add_table(rows=num_rows, cols=2)
cert_table.allow_autofit = False
cert_table.columns[0].width = Inches(3.75)
cert_table.columns[1].width = Inches(3.75)

# Populate table with certifications
for i, cert in enumerate(all_certs):
    row = i // 2
    col = i % 2
    cell = cert_table.cell(row, col)
    cell.width = Inches(3.75)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    # Add bullet character
    run = p.add_run(f"• {cert}")
    run.font.size = Pt(8)

# Save document with versioning
output_dir = "generated"
# Create filename from name (replace spaces with underscores, remove special chars)
name_for_file = personal_info['name'].replace(' ', '_').title()
base_filename = f"{name_for_file}_Resume"

# Create the generate folder if it doesn't exist
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# Find the next available version number
version = 1
while True:
    filename = f"{base_filename}_{version}.docx"
    filepath = os.path.join(output_dir, filename)
    if not os.path.exists(filepath):
        break
    version += 1

# Save the document
doc.save(filepath)
print(f"Resume generated successfully: {filepath}")